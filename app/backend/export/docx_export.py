"""Exportação do Relatório Analítico de Área para .docx (python-docx).

Estampa "RASCUNHO — validação humana" no topo. Recebe um models.Relatorio já montado
(determinístico + o que houver de IA) e devolve os bytes do .docx.
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from ..report import models as M

AZUL = RGBColor(0x16, 0x42, 0x5B)
AMBAR = RGBColor(0x8A, 0x6D, 0x1D)


def _heading(doc, txt: str, size: int = 13):
    p = doc.add_paragraph()
    run = p.add_run(txt)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = AZUL
    return p


def _kv(doc, pairs):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for k, v in pairs:
        cells = t.add_row().cells
        cells[0].text = str(k)
        cells[1].text = "" if v is None else str(v)
    return t


def _table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = "" if v is None else str(v)
    return t


def gerar_docx(rel: M.Relatorio) -> bytes:
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("RELATÓRIO ANALÍTICO DE ÁREA")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = AZUL
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Subsídio para Reunião de CompStat").italic = True

    if rel.rascunho:
        w = doc.add_paragraph()
        wr = w.add_run("RASCUNHO — requer validação humana. A decisão final é do gestor.")
        wr.bold = True
        wr.font.color.rgb = AMBAR

    # 1. Identificação
    _heading(doc, "1. Identificação da Área")
    ident = rel.identificacao
    _kv(doc, [
        ("Área FM", f"{ident.areaFM} — {ident.nomeArea}"),
        ("AISP", ident.aisp), ("DP", ident.dp), ("BPM", ident.bpm),
        ("Base FM", ident.baseFM), ("Subprefeitura", ident.subprefeitura),
        ("Bairros", ", ".join(ident.bairros) or "—"),
        ("Influência de grupo criminoso", ", ".join(ident.influenciaGrupoCriminoso) or "—"),
        ("Período de análise", f"{rel.periodo.de} a {rel.periodo.ate}"),
    ])

    # 2. Resumo Executivo (4 perguntas norteadoras)
    _heading(doc, "2. Resumo Executivo — Perguntas Norteadoras")
    for q in rel.resumoExecutivo:
        p = doc.add_paragraph()
        p.add_run(q.pergunta + " ").bold = True
        diag = q.diagnostico.text if q.diagnostico and q.diagnostico.text else "[a gerar]"
        p.add_run(diag)

    # 3. Ocorrências
    _heading(doc, "3. Ocorrências")
    ind = rel.ocorrencias.indicadores
    _kv(doc, [
        ("Roubos", ind.roubos),
        ("Furtos", "N/A" if ind.furtos is None else ind.furtos),
        ("Total", ind.total),
        ("Ranking entre áreas", ind.rankingEntreAreas),
    ])
    _table(doc, ["Ranking", "Tipo", "Qtd"],
           [(d.rank, d.tipo, d.qtd) for d in rel.ocorrencias.distribuicao])

    # 4. Análise Temporal (texto)
    _heading(doc, "4. Análise Temporal")
    doc.add_paragraph(rel.temporalResumo.text or "[Síntese temporal a gerar.]")

    # 5. Dinâmica Criminal
    _heading(doc, "5. Dinâmica Criminal")
    doc.add_paragraph(rel.dinamicaCriminal.text or "[Dinâmica qualitativa a gerar (depende da extração).]")

    # 6. Efetivo FM
    _heading(doc, "6. Efetivo Empregado — Força Municipal")
    _table(doc, ["Campo", "Situação Atual", "Sugestão", "Justificativa"],
           [(r.campo, r.situacaoAtual, r.sugestao, r.justificativa) for r in rel.efetivoFM])

    # 7. Fatores de incidência
    _heading(doc, "7. Fatores de Incidência Criminal")
    _table(doc, ["Fator", "Órgão Responsável", "Qtd"],
           [(f.fator, f.orgaoResponsavel or "—", f.qtd) for f in rel.fatores.rows])

    # 8. Painel de Coincidências (match)
    _heading(doc, "8. Painel de Coincidências")
    if rel.coincidencias.resumo:
        doc.add_paragraph(rel.coincidencias.resumo)
    _table(doc, ["Score", "Camadas", "Justificativa"],
           [(f"{c.score:.1f}", ", ".join(c.camadas), c.justificativa)
            for c in rel.coincidencias.coincidencias])

    # 9. Plano de Ação
    _heading(doc, "9. Plano de Ação e Responsabilização")
    _table(doc, ["Ação acordada", "Responsável", "Prazo", "Status"],
           [(a.acao, a.responsavel or "—", a.prazo, a.status) for a in rel.planoAcao])

    foot = doc.add_paragraph()
    fr = foot.add_run("CompStat Rio — rascunho automático. Priorização e decisão final são humanas.")
    fr.italic = True
    fr.font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
